# coding=UTF-8

from docx import Document
import pymysql.cursors

"""
This is a tool to generate data dictionary for MySQL table.
This is only yested with MySQL 5.7, but should be compatiable with lower versions too.

Usage: edit the following constants and run this script. Have fun!
"""

# the database to export as data dictionary
DB_NAME = 'test'
DB_HOST = '127.0.0.1'
DB_USER = 'root'
DB_PASSWORD = 'passw0rd'

# the file name of data dictionary
DOC_NAME = 'data_dict.docx'


def update_doc_for_table(doc, table_name, col_list):
    document.add_heading(u'Table Name：' + table_name, level=1)
    document.add_paragraph(u'Description：')

    table = document.add_table(rows=1, cols=4)

    # Set table titles
    cells = table.rows[0].cells
    for i, title in enumerate(tb_titles):
        cells[i].text = title

    # Set table rows
    for col_info in col_list:
        assert len(col_info) == 4, "Invalid format of column info: " + repr(col_info)

        row_cells = table.add_row().cells
        for i, val in enumerate(col_info):
            try:
                row_cells[i].text = val
            except ValueError as e:
                print "Warning: {}, val: {}".format(e, val)
                row_cells[i].text = ''


    document.add_page_break()


def list_tables_for_db(conn, db):
    sql = "select `TABLE_NAME` from `TABLES` where `TABLE_SCHEMA`='{}';".format(db);

    with conn.cursor() as cursor:
        cursor.execute(sql)
        ret = cursor.fetchall()

    table_names = [row[0] for row in ret]
    return table_names


def read_cols_for_table(conn, db, table_name):
    sql = "SELECT `C`.COLUMN_NAME, `C`.COLUMN_TYPE, `C`.IS_NULLABLE, `C`.COLUMN_COMMENT "\
          "FROM `information_schema`.`COLUMNS` `C` "\
          "INNER JOIN `information_schema`.`TABLES` `T` "\
            "ON (`C`.TABLE_SCHEMA = `T`.TABLE_SCHEMA AND `C`.TABLE_NAME = `T`.TABLE_NAME) "\
          "WHERE T.TABLE_SCHEMA='{}' AND T.TABLE_NAME='{}';".format(db, table_name)

    with conn.cursor() as cursor:
        cursor.execute(sql)
        return cursor.fetchall()


if __name__ == '__main__':

    document = Document()
    tb_titles = [u'Column Name', u'Column Type', u'Is Nullable', u'Comment']

    conn = pymysql.connect(host=DB_HOST,
                           user=DB_USER,
                           password=DB_PASSWORD,
                           db='information_schema')

    tables = list_tables_for_db(conn, DB_NAME)
    for tb in tables:
        cols = read_cols_for_table(conn, DB_NAME, tb)
        update_doc_for_table(document, tb, cols)

    document.save(DOC_NAME)
